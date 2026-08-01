import streamlit as st
import google.genai as genai
from google.genai import types
import os
from pathlib import Path
import json
import re
import io
import random
import base64
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend for server
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import matplotlib.ticker as ticker
import numpy as np
from fpdf import FPDF

# Helper: Clean text for export (removes Hangul and Emojis)
def clean_text_for_export(text):
    if not text: return ""
    # Export fonts (standard PDF/Word) handle Latin-1 best. 
    # Encoding to latin-1 with 'ignore' strips Hangul and Emojis.
    return text.encode('latin-1', 'ignore').decode('latin-1')

# Load environment variables (override=True ensures .env wins over stale OS env)
load_dotenv(override=True)

# Page Configuration
st.set_page_config(
    page_title="Elite Prep | SAT Analysis & Adaptive Practice",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Brand Colors & Styles
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    
    html, body, [class*="css"]  {
        font-family: 'Inter', sans-serif;
        color: #1f2937;
        background-color: #f3f4f6;
    }
    
    /* Main Container */
    .block-container {
        padding-top: 2rem;
        padding-bottom: 5rem;
        max-width: 1000px;
    }

    /* Header Styling */
    .main-header {
        font-family: 'Inter', sans-serif;
        color: #0C1E41; /* Elite Navy */
        text-align: center;
        font-weight: 800;
        font-size: 2.2rem;
        margin-bottom: 0.5rem;
        letter-spacing: -0.02em;
    }
    .sub-header {
        color: #6b7280;
        text-align: center;
        margin-bottom: 3rem;
        font-size: 1.1rem;
        font-weight: 400;
    }
    
    /* Card Styling */
    .stCard {
        background-color: white;
        padding: 2rem;
        border-radius: 12px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        margin-bottom: 2rem;
        border: 1px solid #e5e7eb;
    }

    /* Button Styling */
    div.stButton > button {
        background-color: #0C1E41;
        color: white;
        font-weight: 600;
        border-radius: 8px;
        padding: 0.75rem 1.5rem;
        border: none;
        transition: all 0.2s ease;
        width: 100%;
        box-shadow: 0 4px 6px rgba(12, 30, 65, 0.2);
    }
    div.stButton > button:hover {
        background-color: #1a3a6e;
        transform: translateY(-1px);
        box-shadow: 0 6px 8px rgba(12, 30, 65, 0.3);
    }

    /* Download Button (top-right "Download for Word") */
    div.stDownloadButton > button {
        background: linear-gradient(135deg, #0C1E41 0%, #1E4079 100%);
        color: #ffffff !important;
        font-weight: 600;
        font-size: 0.92rem;
        letter-spacing: 0.01em;
        white-space: nowrap;
        border-radius: 10px;
        padding: 0.6rem 1.4rem;
        border: none;
        width: 100%;
        transition: all 0.2s ease;
        box-shadow: 0 4px 12px rgba(12, 30, 65, 0.25);
    }
    div.stDownloadButton > button:hover {
        background: linear-gradient(135deg, #14315B 0%, #2A5599 100%);
        transform: translateY(-2px);
        box-shadow: 0 8px 18px rgba(12, 30, 65, 0.38);
    }
    div.stDownloadButton > button:active {
        transform: translateY(0);
        box-shadow: 0 3px 8px rgba(12, 30, 65, 0.3);
    }
    div.stDownloadButton > button p,
    div.stDownloadButton > button div,
    div.stDownloadButton > button span {
        color: #ffffff !important;
        font-weight: 600 !important;
    }

    /* Practice Topic Buttons (Secondary) */
    .topic-btn {
        border: 1px solid #0C1E41 !important;
        background-color: white !important;
        color: #0C1E41 !important;
    }
    
    /* Status Indicators */
    .status-success { color: #166534; font-weight: 500;}
    .status-error { color: #991b1b; font-weight: 500;}

</style>
""", unsafe_allow_html=True)

# Helper: Get configured Gemini client
def _get_client():
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key and 'api_key' in st.session_state and st.session_state.api_key:
        api_key = st.session_state.api_key
    if api_key:
        return genai.Client(api_key=api_key)
    return None

# Helper: Upload file to Gemini (File API)
def upload_to_gemini(path, mime_type=None):
    try:
        client = _get_client()
        if not client:
            return None
        config = types.UploadFileConfig(mime_type=mime_type) if mime_type else None
        file = client.files.upload(path=path, config=config)
        return file
    except Exception as e:
        print(f"Error uploading {path}: {e}")
        return None

# Helper: Load Gemini Model
def get_gemini_response(input_prompt, content_parts, temperature=0.2):
    try:
        client = _get_client()
        if not client:
            st.error("Google API Key not found.")
            return None

        full_payload = [input_prompt] + content_parts

        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=full_payload,
            config=types.GenerateContentConfig(
                temperature=temperature,
                max_output_tokens=32768,
            ),
        )

        # Check finish_reason to detect truncation
        finish_reason = None
        if response.candidates:
            finish_reason = response.candidates[0].finish_reason
            # finish_reason == 2 → MAX_TOKENS (truncated)
            if finish_reason == 2:
                st.warning("⚠️ The report was truncated due to length limits. Showing partial output. Consider simplifying the request.")

        try:
            return response.text
        except ValueError:
            st.error(f"Generation stopped. Finish Reason: {finish_reason}")
            if finish_reason == 2 and response.candidates and response.candidates[0].content.parts:
                return response.candidates[0].content.parts[0].text
            return None
    except Exception as e:
        st.error(f"Error calling Gemini API: {str(e)}")
        return None

# Helper: Extract JSON from text
def extract_json_topics(text):
    try:
        # Find JSON block using regex
        match = re.search(r'```json\s*(\{.*?\})\s*```', text, re.DOTALL)
        if match:
            return json.loads(match.group(1))
        else:
            # Try finding just brace content if no markdown tags
            match = re.search(r'\{.*\}', text, re.DOTALL)
            if match:
                return json.loads(match.group(0))
    except Exception as e:
        print(f"JSON Parse Error: {e}")
    return None

# Helper: Force Formatting for Options
def ensure_formatting(text):
    if not text: return ""
    # Remove <br> tags
    text = re.sub(r'<br\s*/?>', '', text)
    # Regex to find A) B) C) D) that are NOT at start of line
    pattern = r"(\s)([A-D]\)) "
    formatted_text = re.sub(pattern, r"\n\n\2 ", text)
    return formatted_text

# Helper: Execute matplotlib code and return PNG image bytes
def execute_figure_code(code_str):
    """Execute matplotlib code string and return PNG image as bytes."""
    try:
        plt.close('all')  # Close any previous figures
        
        # Pre-process the code to fix common issues
        lines = code_str.split('\n')
        cleaned_lines = []
        for line in lines:
            stripped = line.strip()
            # Skip import statements (plt, np, io, math already available)
            if stripped.startswith('import ') or stripped.startswith('from '):
                continue
            # Replace plt.show() with savefig
            if stripped == 'plt.show()':
                continue
            cleaned_lines.append(line)
        
        cleaned_code = '\n'.join(cleaned_lines)
        
        # If no savefig call exists, append one
        if 'savefig' not in cleaned_code:
            cleaned_code += "\nplt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')"
        
        # Create a BytesIO buffer for the image
        buf = io.BytesIO()
        
        # Set up execution namespace with common imports available
        exec_namespace = {
            'plt': plt,
            'np': np,
            'io': io,
            'buf': buf,
            'math': __import__('math'),
            'matplotlib': matplotlib,
            'patches': patches,
            'ticker': ticker,
        }
        
        # Execute the code
        exec(cleaned_code, exec_namespace)
        
        # If buf is still empty, try to save current figure
        if buf.tell() == 0:
            fig = plt.gcf()
            if fig.get_axes():
                fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                           facecolor='white', edgecolor='none')
        
        buf.seek(0)
        img_bytes = buf.getvalue()
        plt.close('all')
        
        if len(img_bytes) > 0:
            return img_bytes
        return None
    except Exception as e:
        plt.close('all')
        # Silence warning as requested by user
        # st.warning(f"⚠️ Figure generation error: {e}")
        return None

# Helper: Parse Gemini response into text and figure segments
def parse_response_with_figures(response_text):
    """
    Parse response containing ```python-figure code blocks.
    Returns list of segments: [{"type":"text","content":"..."}, {"type":"figure","image":bytes,"code":"..."}]
    """
    if not response_text:
        return [{"type": "text", "content": ""}]
    
    # Pattern to match ```python-figure ... ``` OR ```python ... ``` blocks containing plt
    pattern = r'```(?:python-figure|python)\s*\n(.*?)```'
    
    segments = []
    last_end = 0
    
    for match in re.finditer(pattern, response_text, re.DOTALL):
        # Add text before this code block
        text_before = response_text[last_end:match.start()].strip()
        if text_before:
            segments.append({"type": "text", "content": text_before})
        
        # Execute the figure code
        code = match.group(1).strip()
        img_bytes = execute_figure_code(code)
        
        if img_bytes:
            segments.append({"type": "figure", "image": img_bytes, "code": code})
        else:
            # If figure generation failed, add a note
            segments.append({"type": "text", "content": "*[Figure could not be generated]*"})
        
        last_end = match.end()
    
    # Add remaining text after last code block
    remaining = response_text[last_end:].strip()
    if remaining:
        segments.append({"type": "text", "content": remaining})
    
    # If no figures found, return entire text as single segment
    if not segments:
        segments.append({"type": "text", "content": response_text})
    
    return segments

# Helper: Convert LaTeX math to readable plain text
def _latex_to_readable(text):
    """Convert LaTeX math expressions to readable Unicode text."""
    if not text:
        return text
    
    def process_math(match):
        math = match.group(1)
        # Handle nested fractions (inner to outer)
        max_iter = 10
        while r'\frac' in math and max_iter > 0:
            math = re.sub(r'\\frac\{([^{}]*)\}\{([^{}]*)\}', r'(\1)/(\2)', math)
            max_iter -= 1
        # Common LaTeX symbols
        replacements = {
            r'\times': '×', r'\div': '÷', r'\pm': '±',
            r'\leq': '≤', r'\geq': '≥', r'\neq': '≠',
            r'\rightarrow': '→', r'\leftarrow': '←',
            r'\cdot': '·', r'\approx': '≈', r'\infty': '∞',
            r'\pi': 'π', r'\theta': 'θ', r'\alpha': 'α', r'\beta': 'β',
            r'\sqrt': '√', r'\le': '≤', r'\ge': '≥',
        }
        for latex_cmd, symbol in replacements.items():
            math = math.replace(latex_cmd, symbol)
        # Subscripts: Try common ones first
        subscripts = {'0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄', '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉'}
        for k, v in subscripts.items():
            math = math.replace('_{' + k + '}', v).replace('_' + k, v)
        
        # Superscripts: Handle common ones (Latin-1 compatible for PDF)
        superscripts = {'1': '¹', '2': '²', '3': '³'}
        for k, v in superscripts.items():
            math = math.replace('^{' + k + '}', v).replace('^' + k, v)
            
        # Fallback for other superscripts/subscripts
        math = re.sub(r'_\{([^}]*)\}', r'_\1', math)
        math = re.sub(r'\^\{([^}]*)\}', r'^(\1)', math)
        
        # Remove remaining backslashes from unknown commands
        math = re.sub(r'\\([a-zA-Z]+)\s*', r'\1 ', math)
        # Clean extra spaces
        math = re.sub(r'\s+', ' ', math).strip()
        return math
    
    # Process $$...$$ blocks first, then $...$
    text = re.sub(r'\$\$\s*([^$]+?)\s*\$\$', process_math, text)
    text = re.sub(r'\$([^$]+?)\$', process_math, text)
    
    # Also convert bare LaTeX commands outside $ delimiters
    while r'\frac' in text:
        prev = text
        text = re.sub(r'\\frac\{([^{}]*)\}\{([^{}]*)\}', r'(\1)/(\2)', text)
        if text == prev:
            break
    
    bare_replacements = {
        r'\times': '×', r'\cdot': '·', r'\rightarrow': '→',
        r'\approx': '≈', r'\leq': '≤', r'\geq': '≥', r'\neq': '≠',
        r'\pm': '±', r'\div': '÷',
    }
    for latex_cmd, symbol in bare_replacements.items():
        text = text.replace(latex_cmd, symbol)
    
    # Convert bare superscripts for export
    text = text.replace('^{2}', '²').replace('^2', '²')
    text = text.replace('^{3}', '³').replace('^3', '³')
    text = text.replace('^{1}', '¹').replace('^1', '¹')
    
    return text

# Helper: Convert Markdown to Word (.docx) document
def convert_markdown_to_docx(md_text):
    if not md_text: return None
    
    doc = Document()
    
    # Set default font for Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    
    # Configure heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Arial'
        heading_style.font.color.rgb = RGBColor(0x0C, 0x1E, 0x41)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    code_block_lang = ""
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Handle code block start/end
        if stripped.startswith('```'):
            if not in_code_block:
                # Starting a code block
                code_block_lang = stripped[3:].strip().lower()
                in_code_block = True
                code_block_lines = []
                i += 1
                continue
            else:
                # Ending a code block
                in_code_block = False
                if code_block_lang in ('python-figure', 'python'):
                    # Execute matplotlib code and embed image
                    code_str = '\n'.join(code_block_lines)
                    img_bytes = execute_figure_code(code_str)
                    if img_bytes:
                        img_buf = io.BytesIO(img_bytes)
                        doc.add_picture(img_buf, width=Inches(4.5))
                        # Center the image
                        last_para = doc.paragraphs[-1]
                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # Regular code block - add as formatted text
                    code_text = '\n'.join(code_block_lines)
                    if code_text.strip():
                        para = doc.add_paragraph()
                        run = para.add_run(code_text)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                code_block_lang = ""
                code_block_lines = []
                i += 1
                continue
        
        if in_code_block:
            code_block_lines.append(line.rstrip())
            i += 1
            continue
        
        # Strip blockquote markers: "> text" → "text", ">" → ""
        while stripped.startswith('>'):
            stripped = stripped[1:].strip()
        
        # Empty line → skip (no blank rows)
        if not stripped:
            i += 1
            continue
        
        # Convert LaTeX in the line
        stripped = _latex_to_readable(stripped)
        
        # Horizontal rule / Page Break
        if stripped in ('---', '--- PAGE BREAK ---') or stripped.startswith('─'):
            para = doc.add_paragraph()
            run = para.add_run('━' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue
        
        # Headers (### > ## > #)
        if stripped.startswith('### '):
            heading = doc.add_heading(level=3)
            _add_formatted_text(heading, stripped[4:])
            i += 1
            continue
        elif stripped.startswith('## '):
            heading = doc.add_heading(level=2)
            _add_formatted_text(heading, stripped[3:])
            i += 1
            continue
        elif stripped.startswith('# '):
            heading = doc.add_heading(level=1)
            _add_formatted_text(heading, stripped[2:])
            i += 1
            continue
        
        # Bullet points: * or - at start
        if re.match(r'^[\*\-]\s+', stripped):
            content = re.sub(r'^[\*\-]\s+', '', stripped)
            para = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(para, content)
            i += 1
            continue
        
        # Numbered items: "1." or "1)" at start — render as text to preserve original numbering
        num_match = re.match(r'^(\d+)([.)\t])\s*(.*)', stripped)
        if num_match:
            num = num_match.group(1)
            content = num_match.group(3)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.2)
            run = para.add_run(f"{num}. ")
            run.bold = True
            _add_formatted_text(para, content)
            i += 1
            continue
        
        # Answer option lines: A) B) C) D) — with slight indent
        option_match = re.match(r'^([A-D])\)\s+(.*)', stripped)
        if option_match:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(f"{option_match.group(1)})  ")
            run.bold = True
            _add_formatted_text(para, option_match.group(2))
            i += 1
            continue
        
        # Regular paragraph
        para = doc.add_paragraph()
        _add_formatted_text(para, stripped)
        i += 1
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def _add_formatted_text(paragraph, text):
    """Add text to a paragraph with bold and italic formatting support."""
    if not text:
        return
    
    # Remove Hangul/Emojis for Word consistency
    text = clean_text_for_export(text)
    
    # First convert any remaining LaTeX
    text = _latex_to_readable(text)
    
    # Split by bold markers **...**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            # Handle superscripts inside bold
            _add_superscripts(paragraph, content, True)
        else:
            # Handle italic *...*
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ip in italic_parts:
                if not ip:
                    continue
                if ip.startswith('*') and ip.endswith('*') and len(ip) > 2:
                    content = ip[1:-1]
                    _add_superscripts(paragraph, content, False, True)
                else:
                    _add_superscripts(paragraph, ip, False, False)

def _add_superscripts(paragraph, text, bold=False, italic=False):
    """Helper to add text with real superscript formatting for Word."""
    # Replace unicode superscripts with ^ marker for parsing
    text = text.replace('²', '^2').replace('³', '^3').replace('¹', '^1')
    text = text.replace('₀', '_0').replace('₁', '_1').replace('₂', '_2').replace('₃', '_3')
    
    # Split by ^ and _
    parts = re.split(r'(\^[0-9a-zA-Z]+|_[0-9a-zA-Z]+)', text)
    for part in parts:
        if part.startswith('^'):
            run = paragraph.add_run(part[1:])
            run.font.superscript = True
            run.bold = bold
            run.italic = italic
        elif part.startswith('_'):
            run = paragraph.add_run(part[1:])
            run.font.subscript = True
            run.bold = bold
            run.italic = italic
        else:
            run = paragraph.add_run(part)
            run.bold = bold
            run.italic = italic

# ============================================================
#  STYLED WORD REPORT  — reproduces the designed sample layout
#  (Eng-Sample_SAT_Analysis_Report & Improvement Plan.docx)
# ============================================================

# Brand palette (hex strings for shading, RGBColor for fonts)
_NAVY   = RGBColor(0x0B, 0x25, 0x45)
_SLATE  = RGBColor(0x14, 0x31, 0x5B)
_BLUE   = RGBColor(0x2E, 0x6B, 0xB8)
_BODY   = RGBColor(0x22, 0x30, 0x3F)
_MUTED  = RGBColor(0x5B, 0x67, 0x72)
_RED    = RGBColor(0xB3, 0x26, 0x1E)
_AMBER  = RGBColor(0xD9, 0x77, 0x06)
_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
_HEROSUB = RGBColor(0xDC, 0xE6, 0xF5)
_FILL_LIGHT = 'EEF2F7'
_GOLD = 'C9A227'
_GRID = 'D9DEE7'


# Word handles Unicode natively, so (unlike the latin-1 PDF path) we strip only
# Hangul and emoji while KEEPING symbols like →, –, · that the sample relies on.
_HANGUL_EMOJI = re.compile(
    r'[가-힣ᄀ-ᇿ㄰-㆏ꥠ-꥿ힰ-퟿'
    r'\U0001F000-\U0001FAFF\U00002600-\U000027BF️]'
)


def _clean_word(text):
    if not text:
        return ""
    return _HANGUL_EMOJI.sub('', text)


def _shade_cell(cell, fill_hex):
    """Apply solid background fill to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _set_cell_margins(cell, top=100, left=140, bottom=100, right=140):
    """Set internal cell padding in twips."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement(f'w:{tag}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _set_table_borders(table, edges):
    """edges: dict like {'bottom':('single','8','C9A227'), 'insideH':('none',None,None)}"""
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{name}')
        val, sz, color = edges.get(name, ('none', None, None))
        el.set(qn('w:val'), val)
        if sz:
            el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color if color else 'auto')
        borders.append(el)
    tblPr.append(borders)


def _no_space(paragraph, before=0, after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.15


def _run(paragraph, text, size, color, bold=False, italic=False):
    r = paragraph.add_run(_clean_word(text))
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def _fixed_table(doc, widths_twips):
    """Create a borderless fixed-layout table with the given column widths (twips)."""
    table = doc.add_table(rows=1, cols=len(widths_twips))
    table.autofit = False
    table.allow_autofit = False
    # fixed layout
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(sum(widths_twips)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    grid = table._tbl.find(qn('w:tblGrid'))
    for gc, w in zip(grid.findall(qn('w:gridCol')), widths_twips):
        gc.set(qn('w:w'), str(w))
    for row in table.rows:
        for cell, w in zip(row.cells, widths_twips):
            cell.width = Pt(w / 20.0)
    return table


def _stat_card(cell, label, big, sub):
    """Fill a stat card cell: small label, large number, small sub-line."""
    _shade_cell(cell, _FILL_LIGHT)
    _set_cell_margins(cell, top=120, left=140, bottom=120, right=140)
    p1 = cell.paragraphs[0]
    _no_space(p1, 0, 2)
    _run(p1, (label or '').upper(), 7, _MUTED, bold=True)
    p2 = cell.add_paragraph(); _no_space(p2, 0, 2)
    _run(p2, big or '', 15, _NAVY, bold=True)
    p3 = cell.add_paragraph(); _no_space(p3, 0, 0)
    _run(p3, sub or '', 8, _MUTED)


def _severity_badge(doc, severity):
    """Small colored pill table indicating weakness severity."""
    sev = (severity or '').lower()
    if 'critical' in sev:
        fill, text = 'B3261E', 'CRITICAL WEAKNESS'
    elif 'foundational' in sev:
        fill, text = '2E6BB8', 'FOUNDATIONAL GAP'
    elif 'inconsist' in sev:
        fill, text = 'D97706', 'INCONSISTENCY'
    else:
        fill, text = '5B6772', (severity or 'AREA OF FOCUS').upper()
    # width scales roughly with text length
    w = min(4200, max(1900, 130 * len(text)))
    table = _fixed_table(doc, [w])
    _set_table_borders(table, {})
    cell = table.rows[0].cells[0]
    _shade_cell(cell, fill)
    _set_cell_margins(cell, top=20, left=100, bottom=20, right=100)
    p = cell.paragraphs[0]; _no_space(p, 0, 0)
    _run(p, text, 7.5, _WHITE, bold=True)


def _heading(doc, text, size=13, color=None, before=16, after=7):
    p = doc.add_paragraph(); _no_space(p, before, after)
    _run(p, text, size, color if color else _NAVY, bold=True)
    return p


def _label_value(doc, label, value, justify=False):
    p = doc.add_paragraph(); _no_space(p, 0, 4.5)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, f"{label}  ", 10, _SLATE, bold=True)
    _run(p, value, 10, _BODY)
    return p


SAT_MAX_SCORE = 1600
TARGET_MIN_GAIN = 50
TARGET_MAX_GAIN = 100


def _compute_target_score(peak):
    """Target score = the student's highest total + 50 to 100 points, capped at
    the SAT maximum of 1600. Returns a display string ('980 - 1030' or '1600')."""
    try:
        peak = int(peak)
    except (TypeError, ValueError):
        return ''
    low = min(peak + TARGET_MIN_GAIN, SAT_MAX_SCORE)
    high = min(peak + TARGET_MAX_GAIN, SAT_MAX_SCORE)
    if low >= high:
        return str(SAT_MAX_SCORE)
    return f"{low} - {high}"


def _peak_total_from_report(md):
    """Highest total score across the student's tests. Reads the ```scores``` block
    first, falling back to the 'Test Trajectory:' line. Returns None if unavailable."""
    if not md:
        return None
    totals = []
    sm = re.search(r'```scores\s*(\[.*?\])\s*```', md, re.DOTALL)
    if sm:
        try:
            for e in json.loads(sm.group(1)):
                t = e.get('total')
                if isinstance(t, (int, float)):
                    totals.append(int(t))
        except Exception as ex:
            print(f"scores block parse error (peak): {ex}")
    if not totals:
        m = re.search(r'Test Trajectory:\s*([^\n]+)', md)
        if m:
            totals = [int(s) for s in re.findall(r'(\d{3,4})', m.group(1))]
    totals = [t for t in totals if 200 <= t <= SAT_MAX_SCORE]
    return max(totals) if totals else None


def enforce_target_score(md):
    """Rewrite the report's 'Target Score' line so it always follows the
    highest-score + 50~100 (max 1600) rule. Returns the report unchanged if the
    peak score cannot be determined."""
    if not md:
        return md
    target = _compute_target_score(_peak_total_from_report(md))
    if not target:
        return md
    new_md, n = re.subn(r'(?im)^(\s*\**\s*Target Score\s*\**\s*:\s*\**\s*)(.*)$',
                        lambda mm: f"{mm.group(1)}{target}", md, count=1)
    return new_md if n else md


def _parse_report(md):
    """Parse the generated markdown report into structured data. Defensive: any
    field that cannot be found is left empty rather than raising."""
    data = {
        'student_name': '', 'test_names': [], 'trajectory': [], 'intro': '',
        'perf_trend': '', 'weaknesses': [], 'target_score': '', 'goal': '',
        'total_hours': '', 'split': '', 'frequency': '', 'weeks': [], 'scores': [],
    }
    if not md:
        return data
    text = md

    def _clean(s):
        return re.sub(r'\*+', '', s).strip()

    # Intro line ("Based on ...")
    m = re.search(r'^(Based on the diagnostic.*?)$', text, re.MULTILINE | re.DOTALL)
    if m:
        data['intro'] = _clean(m.group(1).split('\n\n')[0].split('\n#')[0]).strip()

    # Student name
    m = re.search(r'Student Name:\s*\**\s*([^\n]+)', text)
    if m:
        data['student_name'] = _clean(m.group(1))
    if not data['student_name']:
        m = re.search(r'provided for\s+([A-Z][\w.\-]*(?:\s+[A-Z][\w.\-]*)*)\s+across', text)
        if m:
            data['student_name'] = m.group(1).strip()

    # Test trajectory: "930 (DSAT-01-B) -> 890 (DSAT-02-A) -> ..."
    m = re.search(r'Test Trajectory:\s*([^\n]+)', text)
    if m:
        for sc, tn in re.findall(r'(\d{3,4})\s*\(([^)]+)\)', m.group(1)):
            data['trajectory'].append((int(sc), tn.strip()))
    data['test_names'] = [t for _, t in data['trajectory']]

    # Per-test section scores block (for the trajectory chart)
    sm = re.search(r'```scores\s*(\[.*?\])\s*```', text, re.DOTALL)
    if sm:
        try:
            for e in json.loads(sm.group(1)):
                data['scores'].append({
                    'test': str(e.get('test', '')).strip(),
                    'total': e.get('total'),
                    'math': e.get('math'),
                    'rw': e.get('rw'),
                })
        except Exception as ex:
            print(f"scores block parse error: {ex}")

    # Performance trend
    m = re.search(r'Performance Trend:\s*\**\s*(.+?)(?:\n\s*\n|\n#|\Z)', text, re.DOTALL)
    if m:
        data['perf_trend'] = _clean(m.group(1))

    # Weaknesses:  "### 1. Math: Advanced Math (Critical Weakness)"
    weak_pat = re.compile(
        r'^#{2,3}\s*(\d+)\.\s*([^:]+):\s*(.+?)\s*(?:\(([^)]+)\))?\s*$', re.MULTILINE)
    matches = list(weak_pat.finditer(text))
    for idx, wm in enumerate(matches):
        start = wm.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        # stop the block at the next top-level section if one appears sooner
        nxt = re.search(r'\n##\s+', text[start:end])
        block = text[start:start + nxt.start()] if nxt else text[start:end]
        skills = ''
        sm = re.search(r'Specific Skills Lacking:\s*\**\s*(.+?)(?:\n\s*[-*#]|\n\s*\*\*Curriculum|\Z)',
                       block, re.DOTALL)
        if sm:
            skills = _clean(sm.group(1))
        chapters = []
        cm = re.search(r'Curriculum Chapters:?\**(.*)', block, re.DOTALL)
        chap_src = cm.group(1) if cm else block
        for line in chap_src.split('\n'):
            lm = re.match(r'^\s*[-*]\s+(.+)', line)
            if lm:
                chapters.append(_clean(lm.group(1)))
        data['weaknesses'].append({
            'num': wm.group(1), 'subject': _clean(wm.group(2)),
            'category': _clean(wm.group(3)), 'severity': (wm.group(4) or '').strip(),
            'skills': skills, 'chapters': chapters,
        })

    # Logistics fields
    def _field(label):
        mm = re.search(rf'{label}:\s*\**\s*([^\n]+)', text)
        return _clean(mm.group(1)) if mm else ''
    data['target_score'] = _field('Target Score')
    # The target is always derived from the student's own best result
    # (highest total + 50~100, capped at 1600), never left to the model.
    computed_target = _compute_target_score(_peak_total_from_report(text))
    if computed_target:
        data['target_score'] = computed_target
    data['goal'] = _field('Goal Breakdown')
    data['total_hours'] = _field(r'Total Tutoring Hours')
    data['split'] = _field('Split')
    data['frequency'] = _field('Frequency')

    # Week-by-week plan:  "**Week 1 (Theme):**" then "- Math: ..." / "- RW: ..."
    week_pat = re.compile(r'\*{0,2}Week\s+(\d+)\s*\(([^)]*)\)\**\s*:?', re.IGNORECASE)
    wmatches = list(week_pat.finditer(text))
    for idx, wm in enumerate(wmatches):
        start = wm.end()
        end = wmatches[idx + 1].start() if idx + 1 < len(wmatches) else len(text)
        block = text[start:end]
        math_m = re.search(r'Math:\s*\**\s*([^\n]+)', block)
        rw_m = re.search(r'(?:RW|Reading\s*&?\s*Writing|R&W):\s*\**\s*([^\n]+)', block)
        data['weeks'].append({
            'num': wm.group(1), 'theme': _clean(wm.group(2)),
            'math': _clean(math_m.group(1)) if math_m else '',
            'rw': _clean(rw_m.group(1)) if rw_m else '',
        })
    return data


def _build_trajectory_chart(scores, trajectory):
    """Render the SAT Score Trajectory line chart (Total/Math/RW) as PNG bytes,
    styled to match the sample report. Falls back to a Total-only line if section
    scores are unavailable. Returns None on any failure."""
    try:
        if scores and any(s.get('total') is not None for s in scores):
            labels = [s.get('test', '') for s in scores]
            totals = [s.get('total') for s in scores]
            math = [s.get('math') for s in scores]
            rw = [s.get('rw') for s in scores]
        elif trajectory:
            labels = [t for _, t in trajectory]
            totals = [s for s, _ in trajectory]
            math = rw = None
        else:
            return None

        navy, blue, gold, muted, body = '#0B2545', '#2E6BB8', '#C9A227', '#5B6772', '#22303F'
        plt.close('all')
        fig, ax = plt.subplots(figsize=(9.7, 4.5), dpi=150)
        x = list(range(len(labels)))

        if math is not None and any(v is not None for v in math):
            ax.plot(x, math, color=blue, marker='o', markersize=6, linewidth=1.8,
                    label='Math', zorder=3)
        if rw is not None and any(v is not None for v in rw):
            ax.plot(x, rw, color=gold, marker='s', markersize=6, linewidth=1.8,
                    label='Reading & Writing', zorder=3)
        ax.plot(x, totals, color=navy, marker='D', markersize=8, linewidth=2.6,
                markerfacecolor=navy, markeredgecolor='white', label='Total Score', zorder=5)

        # data labels on the Total line
        for xi, tv in zip(x, totals):
            if tv is None:
                continue
            ax.annotate(str(tv), (xi, tv), textcoords='offset points', xytext=(0, 12),
                        ha='center', fontsize=10, fontweight='bold', color=navy, zorder=6)

        # peak marker
        valid = [(xi, tv) for xi, tv in zip(x, totals) if tv is not None]
        if valid:
            px, pv = max(valid, key=lambda p: p[1])
            ax.scatter([px], [pv], s=240, facecolors='none', edgecolors=gold,
                       linewidths=2.2, zorder=7)
            ax.annotate('Peak', (px, pv), textcoords='offset points', xytext=(0, 28),
                        ha='center', fontsize=9, fontweight='bold', color=gold, zorder=7)

        rng = f"{labels[0]}  →  {labels[-1]}" if labels else ''
        ax.set_title(f"SAT Score Trajectory    ·    {rng}", fontsize=13,
                     fontweight='bold', color=navy, pad=28)
        ax.set_ylabel('Score', fontsize=10, color=muted)
        ax.set_xticks(x)
        ax.set_xticklabels(labels, fontsize=9, color=body)
        ax.tick_params(axis='y', labelsize=9, colors=muted)

        allv = [v for v in (list(totals) + list(math or []) + list(rw or [])) if v is not None]
        if allv:
            lo, hi = min(allv), max(allv)
            ax.set_ylim(max(0, lo - 90), hi + 95)

        ax.grid(axis='y', color='#E5E9F0', linewidth=1)
        ax.set_axisbelow(True)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('#D9DEE7')
        ax.spines['bottom'].set_color('#D9DEE7')
        ax.set_facecolor('white')
        fig.patch.set_facecolor('white')

        ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.11), ncol=3,
                  frameon=False, fontsize=9)

        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        plt.close('all')
        buf.seek(0)
        return buf.getvalue()
    except Exception as e:
        print(f"Trajectory chart error: {e}")
        plt.close('all')
        return None


def convert_report_to_styled_docx(md_text):
    """Build a Word document matching the designed sample layout."""
    if not md_text:
        return None
    try:
        d = _parse_report(md_text)
        # If the totals trajectory line was missing, derive it from the scores block
        if not d['trajectory'] and d['scores']:
            d['trajectory'] = [(s['total'], s['test']) for s in d['scores']
                               if s.get('total') is not None]
            d['test_names'] = [t for _, t in d['trajectory']]
        doc = Document()
        # base font
        normal = doc.styles['Normal'].font
        normal.name = 'Calibri'
        normal.size = Pt(10)
        normal.color.rgb = _BODY
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # ---- Hero banner ----
        hero = _fixed_table(doc, [10080])
        _set_table_borders(hero, {})
        hc = hero.rows[0].cells[0]
        _shade_cell(hc, '3D6DB0')  # lighter navy-blue (was 0B2545)
        _set_cell_margins(hc, top=240, left=300, bottom=240, right=300)
        hp = hc.paragraphs[0]; _no_space(hp, 0, 2)
        _run(hp, 'ELITE PREP', 15, _WHITE, bold=True)
        hp2 = hc.add_paragraph(); _no_space(hp2, 0, 0)
        _run(hp2, 'SAT Analysis Report & Improvement Plan', 11, _HEROSUB, bold=True)

        # ---- Info bar (student | test range) ----
        info = _fixed_table(doc, [5040, 5040])
        _set_table_borders(info, {'bottom': ('single', '8', _GOLD)})
        c_left, c_right = info.rows[0].cells
        _set_cell_margins(c_left, top=80, left=0, bottom=80, right=0)
        _set_cell_margins(c_right, top=80, left=0, bottom=80, right=0)
        pl = c_left.paragraphs[0]; _no_space(pl, 0, 0)
        _run(pl, 'STUDENT  ', 7.5, _MUTED, bold=True)
        _run(pl, d['student_name'] or 'Student', 11, _NAVY, bold=True)
        pr = c_right.paragraphs[0]; _no_space(pr, 0, 0)
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if d['test_names']:
            rng = f"{len(d['test_names'])} Diagnostic Tests  ·  {d['test_names'][0]} → {d['test_names'][-1]}"
        else:
            rng = 'Diagnostic Report'
        _run(pr, rng, 9, _MUTED)

        doc.add_paragraph()  # spacer

        # ---- Score Trajectory Overview: chart + subtitle + stat cards ----
        totals = [s for s, _ in d['trajectory']]
        if totals:
            _heading(doc, 'Score Trajectory Overview', size=11, before=4, after=3)

            # Trajectory line chart (Total / Math / RW)
            chart_png = _build_trajectory_chart(d['scores'], d['trajectory'])
            if chart_png:
                pic_p = doc.add_paragraph(); _no_space(pic_p, 0, 4)
                pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_p.add_run().add_picture(io.BytesIO(chart_png), width=Inches(6.9))

            peak = max(totals); peak_test = d['trajectory'][totals.index(peak)][1]
            latest, latest_test = d['trajectory'][-1]
            lo, hi = min(totals), max(totals)
            sub = doc.add_paragraph(); _no_space(sub, 0, 6); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(sub,
                 f"Total score varies across the {len(totals)} diagnostics, peaking at {peak} "
                 f"on {peak_test} and finishing at {latest} on {latest_test}; "
                 f"Math and Reading & Writing shown for context.",
                 8, _MUTED, italic=True)
            cards = _fixed_table(doc, [3200, 3200, 3200])
            _set_table_borders(cards, {'insideH': ('single', '24', 'FFFFFF'),
                                       'insideV': ('single', '24', 'FFFFFF')})
            _stat_card(cards.rows[0].cells[0], 'SCORE RANGE', f"{lo} – {hi}",
                       f"Total, across {len(totals)} tests")
            _stat_card(cards.rows[0].cells[1], 'PEAK SCORE', str(peak), peak_test)
            _stat_card(cards.rows[0].cells[2], 'LATEST SCORE', str(latest), latest_test)
            doc.add_paragraph()

        # ---- Intro paragraph ----
        if d['intro']:
            ip = doc.add_paragraph(); _no_space(ip, 0, 2); ip.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _run(ip, d['intro'], 10, _BODY)

        # ---- Diagnostic Report Overview ----
        _heading(doc, 'Diagnostic Report Overview')
        if d['student_name']:
            _label_value(doc, 'Student Name:', d['student_name'])
        if d['trajectory']:
            traj = '  →  '.join(f"{s} ({t})" for s, t in d['trajectory'])
            _label_value(doc, 'Test Trajectory:', traj)
        if d['perf_trend']:
            _label_value(doc, 'Performance Trend:', d['perf_trend'], justify=True)

        # ---- Key Weaknesses ----
        if d['weaknesses']:
            _heading(doc, 'Key Weaknesses & Curricular Solutions')
            for w in d['weaknesses']:
                _severity_badge(doc, w['severity'])
                tp = doc.add_paragraph(); _no_space(tp, 4, 4)
                _run(tp, f"{w['num']}. ", 11, _BLUE, bold=True)
                _run(tp, f"{w['subject']}: {w['category']}", 11, _SLATE, bold=True)
                if w['skills']:
                    sp = doc.add_paragraph(); _no_space(sp, 0, 5.5)
                    sp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    _run(sp, 'Specific Skills Lacking:  ', 10, _SLATE, bold=True)
                    _run(sp, w['skills'], 10, _BODY)
                if w['chapters']:
                    cp = doc.add_paragraph(); _no_space(cp, 0, 2.5)
                    _run(cp, 'Curriculum Chapters', 9.5, _BLUE, bold=True)
                    for ch in w['chapters']:
                        bp = doc.add_paragraph(style='List Bullet'); _no_space(bp, 0, 2)
                        _run(bp, ch, 9.5, _BODY)

        # ---- Tutoring Recommendations & Logistics ----
        if any([d['target_score'], d['total_hours'], d['frequency'], d['goal'], d['split']]):
            _heading(doc, 'Tutoring Recommendations & Logistics')
            # derive sub-lines
            def _grab(pattern, src):
                mm = re.search(pattern, src or '')
                return mm.group(0) if mm else ''
            target_sub = ''
            gm = re.findall(r'(\d{3,4}\s*[–\-]\s*\d{3,4})', d['goal'])
            if len(gm) >= 2:
                target_sub = f"Math {gm[0]} · RW {gm[1]}"
            hours_sub = ''
            hm = re.search(r'(\d+)\s*Hours?\s*Math\s*/\s*(\d+)\s*Hours?\s*Reading', d['split'])
            if hm:
                hours_sub = f"{hm.group(1)} Math / {hm.group(2)} RW"
            if any([d['target_score'], d['total_hours'], d['frequency']]):
                logi = _fixed_table(doc, [3200, 3200, 3200])
                _set_table_borders(logi, {'insideH': ('single', '24', 'FFFFFF'),
                                          'insideV': ('single', '24', 'FFFFFF')})
                _stat_card(logi.rows[0].cells[0], 'TARGET SCORE', d['target_score'] or '—', target_sub)
                _stat_card(logi.rows[0].cells[1], 'TOTAL HOURS', d['total_hours'] or '—', hours_sub)
                _stat_card(logi.rows[0].cells[2], 'FREQUENCY', d['frequency'] or '—', '')
                doc.add_paragraph()
            if d['goal']:
                _label_value(doc, 'Goal Breakdown:', d['goal'], justify=True)
            if d['split']:
                _label_value(doc, 'Hours Split:', d['split'], justify=True)

        # ---- Week-by-Week Study Plan ----
        if d['weeks']:
            _heading(doc, 'Week-by-Week Study Plan')
            wt = _fixed_table(doc, [1250, 2650, 5700])
            _set_table_borders(wt, {'top': ('single', '4', '0B2545'),
                                    'bottom': ('single', '4', '0B2545'),
                                    'insideH': ('single', '2', _GRID),
                                    'insideV': ('single', '2', _GRID)})
            hdr = wt.rows[0].cells
            for cell, txt in zip(hdr, ('Week', 'Focus', 'Curriculum')):
                _shade_cell(cell, '0B2545')
                _set_cell_margins(cell, top=90, left=130, bottom=90, right=130)
                p = cell.paragraphs[0]; _no_space(p, 0, 0)
                _run(p, txt, 9, _WHITE, bold=True)
            for i, wk in enumerate(d['weeks']):
                row = wt.add_row().cells
                fill = _FILL_LIGHT if i % 2 else 'FFFFFF'
                for cell in row:
                    _shade_cell(cell, fill)
                    _set_cell_margins(cell, top=90, left=130, bottom=90, right=130)
                p0 = row[0].paragraphs[0]; _no_space(p0, 0, 0)
                _run(p0, f"Week {wk['num']}", 9, _NAVY, bold=True)
                p1 = row[1].paragraphs[0]; _no_space(p1, 0, 0)
                _run(p1, wk['theme'], 8.5, _BODY)
                p2 = row[2].paragraphs[0]; _no_space(p2, 0, 0)
                if wk['math']:
                    _run(p2, 'Math: ', 8.5, _BLUE, bold=True)
                    _run(p2, wk['math'], 8.5, _BODY)
                if wk['rw']:
                    p2b = row[2].add_paragraph(); _no_space(p2b, 1, 0)
                    _run(p2b, 'RW: ', 8.5, _BLUE, bold=True)
                    _run(p2b, wk['rw'], 8.5, _BODY)

        # ---- Footer line ----
        doc.add_paragraph()
        fp = doc.add_paragraph(); _no_space(fp, 6, 0); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(fp, 'Elite Prep  ·  SAT Analysis Report & Improvement Plan', 7.5, _MUTED)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        print(f"Styled DOCX generation error: {e}")
        # Fall back to the plain markdown conversion so the button still works
        return convert_markdown_to_docx(md_text)


class ElitePDF(FPDF):
    def header(self):
        # Modern Header with Navy bar
        self.set_fill_color(12, 30, 65) # Elite Navy
        self.rect(0, 0, 210, 15, 'F')
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(255, 255, 255)
        self.set_y(5)
        self.cell(0, 5, "ELITE PREP | SAT ADAPTIVE PRACTICE STATION", align="C", ln=True)
        self.ln(10)

    def footer(self):
        # Subtle Footer
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(100, 100, 100)
        self.cell(0, 10, f"Confidential Elite Prep - Page {self.page_no()}/{{nb}}", align="C")

def convert_markdown_to_pdf(md_text):
    if not md_text: return None
    
    pdf = ElitePDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    
    elite_navy = (12, 30, 65)
    
    # Block grouping to keep questions together
    raw_lines = md_text.split('\n')
    blocks = []
    temp_block = []
    
    for line in raw_lines:
        s = line.strip()
        # Detect new sections or questions to start a new block
        # Match question numbers like 1. or **1.**
        is_q = re.match(r'^(\*\*|)(\d+)[.)]', s)
        if s.startswith('### ') or s.startswith('## ') or s.startswith('# ') or \
           is_q or s.startswith('Question ') or \
           s.startswith('Answer Key') or s == '--- PAGE BREAK ---':
            if temp_block: blocks.append(temp_block)
            temp_block = [line]
        else:
            temp_block.append(line)
    if temp_block: blocks.append(temp_block)
    
    for block in blocks:
        # Estimate height to prevent messy breaks
        est_h = len(block) * 8
        if any('```python-figure' in l for l in block): est_h += 80
        if pdf.get_y() + est_h > (pdf.h - 25): pdf.add_page()
        
        in_code = False
        code_lines = []
        code_lang = ""
        
        for line in block:
            pdf.set_x(18) # Relaxed margins
            s = line.strip()
            
            if s.startswith('```'):
                if not in_code:
                    in_code = True; code_lang = s[3:].strip().lower(); code_lines = []; continue
                else:
                    in_code = False
                    if code_lang in ('python-figure', 'python'):
                        img = execute_figure_code('\n'.join(code_lines))
                        if img: pdf.image(io.BytesIO(img), x=35, w=140); pdf.ln(5)
                    else:
                        pdf.set_font("Courier", size=9); pdf.set_fill_color(245, 245, 245)
                        pdf.multi_cell(0, 5, clean_text_for_export('\n'.join(code_lines)), fill=True)
                        pdf.set_font("Helvetica", size=11)
                    continue
            if in_code: code_lines.append(line); continue
                
            # Process text and clean markdown artifacts
            text = _latex_to_readable(s)
            while text.startswith('>'): text = text[1:].strip()
            
            # Regex for SAT elements
            q_m = re.match(r'^(\*\*|)(\d+)[.)\s\*]+(.*)', text)
            o_m = re.match(r'^(\*\*|)([A-D])\)[.\s\*]*(.*)', text)
            
            if text.startswith('### '):
                pdf.set_text_color(*elite_navy); pdf.set_font("Helvetica", "B", 14)
                pdf.cell(0, 12, clean_text_for_export(text[4:]), ln=True)
                pdf.set_font("Helvetica", size=11); pdf.set_text_color(0)
            elif text.startswith('## ') or text.startswith('# '):
                pdf.set_text_color(*elite_navy); pdf.set_font("Helvetica", "B", 18)
                label = text[text.find(' ')+1:]
                pdf.cell(0, 16, clean_text_for_export(label), ln=True)
                pdf.set_font("Helvetica", size=11); pdf.set_text_color(0)
            elif text in ('---', '--- PAGE BREAK ---'):
                if text == '--- PAGE BREAK ---': pdf.add_page()
                else:
                    pdf.ln(2); pdf.set_draw_color(200, 200, 200)
                    pdf.line(20, pdf.get_y(), 190, pdf.get_y()); pdf.ln(5)
            elif q_m:
                pdf.set_font("Helvetica", "B", 11); pdf.set_text_color(*elite_navy)
                num = q_m.group(2); content = q_m.group(3).strip().replace('**', '')
                pdf.write(7, f"{num}. "); pdf.set_font("Helvetica", size=11); pdf.set_text_color(0)
                pdf.multi_cell(0, 7, clean_text_for_export(content))
            elif o_m:
                pdf.set_x(28); pdf.set_font("Helvetica", "B", 10)
                letter = o_m.group(2); o_text = o_m.group(3).strip().replace('**', '')
                pdf.write(6, f"{letter}) "); pdf.set_font("Helvetica", size=11)
                pdf.multi_cell(0, 6, clean_text_for_export(o_text))
            else:
                cleaned = clean_text_for_export(text).replace('**', '').replace('__', '')
                if cleaned.strip(): pdf.multi_cell(0, 6, cleaned)
                else: pdf.ln(3)

                
    return bytes(pdf.output())

# --- Sidebar ---
with st.sidebar:
    # Logo & Brand (Centered)
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        if Path("Elite-Logo.png").exists():
            st.image("Elite-Logo.png", width=150)
    
    st.markdown("<h2 style='text-align: center; color: #0C1E41; margin-top: -10px;'>Elite Prep</h2>", unsafe_allow_html=True)
    
    st.markdown("---")
    
    # API Key (Hidden Label)
    env_key = os.getenv("GOOGLE_API_KEY")
    if env_key:
        st.success("✅ API Key Loaded")
    else:
        api_key_input = st.text_input("Enter Google API Key", type="password")
        if api_key_input:
            st.session_state.api_key = api_key_input
    
    st.markdown("---")
    st.markdown("### 📂 Upload the Student’s SAT Test Result Report")
    uploaded_files = st.file_uploader(
        "Upload Diagnostic PDF(s)", 
        type=['pdf', 'jpg', 'png'], 
        accept_multiple_files=True,
        label_visibility="collapsed"
    )
    
# Header row: title (left) + top-right "Download for Word" button placeholder
head_title, head_btn = st.columns([4, 1.4], vertical_alignment="center")
with head_title:
    st.markdown("<h1 class='main-header' style='text-align:left;'>Elite Prep | SAT Analysis Report & Improvement Plan</h1>", unsafe_allow_html=True)
with head_btn:
    word_btn_placeholder = st.empty()


# Initialization
if 'analysis_done' not in st.session_state:
    st.session_state.analysis_done = False
if 'full_report' not in st.session_state:
    st.session_state.full_report = ""
if 'weak_topics' not in st.session_state:
    st.session_state.weak_topics = None
if 'practice_result' not in st.session_state:
    st.session_state.practice_result = ""
if 'all_topics' not in st.session_state:
    st.session_state.all_topics = None
if 'manual_practice_result' not in st.session_state:
    st.session_state.manual_practice_result = ""

# ==========================================
# Student Analysis & Adaptive Practice
# ==========================================
with st.container():
    # 1. Report Generation Phase
    if not st.session_state.analysis_done:
        if uploaded_files:
            st.success(f"📂 {len(uploaded_files)} Report(s) Ready.")
            
            if st.button("🚀 Generate Analysis Report"):
                with st.spinner("Analyzing Student Data... (This takes about 10-20 seconds)"):
                    
                    
                    # Load User Files
                    student_parts = []
                    for up_file in uploaded_files:
                        student_parts.append(
                            types.Part.from_bytes(
                                data=up_file.getvalue(),
                                mime_type=up_file.type
                            )
                        )
                    
                    # --- NEW: Load Knowledge Base ---
                    knowledge_base_text = ""
                    kb_path = Path("knowledge/sat_curriculum.json")
                    if kb_path.exists():
                        try:
                            with open(kb_path, "r", encoding="utf-8") as f:
                                knowledge_base_text = f.read()
                        except Exception as e:
                            st.warning(f"Error loading knowledge base: {e}")
                    
                    all_content = student_parts
                    
                    # Report Prompt
                    prompt = f"""
You are Elite Prep's Senior SAT Consultant. You have been provided with:
1. The student's Diagnostic Test Result images (the actual score sheets to analyze).
2. The official SAT Curriculum Knowledge Base (below).

**KNOWLEDGE BASE:**
```json
{knowledge_base_text}
```

**TASK:** Produce a formal diagnostic report. Output 100% in English. No Korean characters anywhere.

---

**OUTPUT FORMAT — follow this structure exactly, word-for-word section titles:**

Based on the diagnostic test results provided for [Student Name] across [N] practice tests ([list test names]), here is a detailed analysis of the student's performance, key weaknesses, and a targeted study plan.

## Diagnostic Report Overview
Student Name: [Full Name]
Test Trajectory: [Score (Test 1)] → [Score (Test 2)] → ... → [Score (Test N)]
Performance Trend: [2–3 sentence narrative. Mention total score improvement, omitted questions trend, Math plateau, and RW consistency. Be specific with numbers.]

---

## Key Weaknesses & Curricular Solutions

[For each weakness, use the structure below. Number them 1, 2, 3, 4. Always cover exactly 2 Math and 2 Reading & Writing weaknesses. DO NOT compare with other students, focus solely on the requested student. Use the KNOWLEDGE BASE to map the weakness to the correct Curricular Category and Chapters.]

### [N]. [Subject]: [Weakness Category Name] ([Severity: e.g., "Critical Weakness" / "Foundational Gap" / "Inconsistency"])
**Specific Skills Lacking:** [2–3 sentence explanation. Reference specific test scores and percentages from the student's actual results. Describe exactly what conceptual knowledge is missing based on the Knowledge Base 'common_mistakes' and 'key_concepts'.]
**Curriculum Chapters:**
- Chapter [X]: [Chapter Name]
- Chapter [X]: [Chapter Name]
- (add more as needed)

---

## Tutoring Recommendations & Logistics
**Target Score:** [MANDATORY RULE — ignore the Knowledge Base score tiers here. Take the HIGHEST total score the student actually earned across all the tests provided, then add 50 to 100 points. Write it as a range, e.g. highest = 1000 -> "1050 - 1100". Never exceed 1600: cap both ends at 1600 (e.g. highest = 1560 -> "1600").]
**Goal Breakdown:** [1–2 sentences: Math target, RW target, and why. The Math and RW targets MUST add up to the Target Score range above, and each section is capped at 800.]
**Total Tutoring Hours:** [N] Hours [use the Scoring Guidelines]
**Split:** [Math Hours] Hours Math / [RW Hours] Hours Reading & Writing. ([1 sentence justification.])
**Frequency:** [Frequency from Scoring Guidelines]

### Week-by-Week Study Plan
**Week 1 ([Theme]):**
- Math: [Chapter(s)]
- RW: [Chapter(s) + brief focus note]

**Week 2 ([Theme]):**
- Math: [Chapter(s)]
- RW: [Chapter(s) + brief focus note]

[Continue for the recommended number of weeks (typically 6-8). Second to last week = timed drills. Last week = full practice test review.]

---

**CRITICAL OUTPUT INSTRUCTION:**
After the full report, append a RAW JSON block listing the weak lesson subjects for the practice question generator. Format exactly:
```json
{{
    "English": ["Chapter X: Name", "Chapter X: Name"],
    "Math": ["Chapter X: Name", "Chapter X: Name"]
}}
```

Then, on a new line, append a RAW scores data block for the Score Trajectory chart. Read each test's Total, Math section, and Reading & Writing section scores directly from the score reports, listed in chronological order (same order as the Test Trajectory). Use integers only. Format exactly:
```scores
[
    {{"test": "DSAT-01-B", "total": 930, "math": 430, "rw": 500}},
    {{"test": "DSAT-02-A", "total": 890, "math": 390, "rw": 500}}
]
```
"""
                    
                    response_text = get_gemini_response(prompt, all_content)
                    
                    if response_text:
                        # Force the target score to follow the highest-score + 50~100
                        # (max 1600) rule so the on-screen and Word reports agree.
                        st.session_state.full_report = enforce_target_score(response_text)
                        st.session_state.analysis_done = True
                        
                        # Parse JSON
                        topics = extract_json_topics(response_text)
                        if topics:
                            st.session_state.weak_topics = topics
                        else:
                            st.warning("Could not auto-detect weak topics for the practice menu. (JSON Parsing Failed)")
                            # Fallback default
                            st.session_state.weak_topics = {"English": ["General Practice"], "Math": ["General Practice"]}
                        
                        st.rerun() # Reload to show report
                    else:
                        st.error("Analysis Failed. Please try again.")
        else:
            st.info("👋 Please upload the student's SAT test result in the sidebar to start the analysis.")

    # 2. Display Report
    if st.session_state.analysis_done:
        st.markdown("### 📊 Analysis Report")

        # Remove the JSON block from display if possible for cleanliness (Optional, keeping it simple for now)
        clean_report = re.sub(r'```json\s*\{.*?\}\s*```', '', st.session_state.full_report, flags=re.DOTALL)
        # Also hide the raw scores data block from the on-screen report
        clean_report = re.sub(r'```scores\s*\[.*?\]\s*```', '', clean_report, flags=re.DOTALL)

        # Download buttons
        from datetime import datetime
        date_str = datetime.now().strftime("%Y%m%d")
        WORD_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # Styled Word document that mirrors the designed sample layout.
        # Pass the RAW report so the builder can read the scores block for the chart.
        styled_docx = convert_report_to_styled_docx(st.session_state.full_report)

        # Top-right header button (the red-circle location) — the only download button
        if styled_docx:
            word_btn_placeholder.download_button(
                label="📥 Download for Word",
                data=styled_docx,
                file_name=f"SAT_Analysis_Report_{date_str}.docx",
                mime=WORD_MIME,
                use_container_width=True,
                key="word_dl_top",
            )

        with st.expander("📄 View Full Analysis Report", expanded=True):
            st.markdown(clean_report)
